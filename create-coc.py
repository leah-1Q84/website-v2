from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page setup - A4
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

PURPLE = RGBColor(0x2A, 0x14, 0x82)
DARK = RGBColor(0x1A, 0x1A, 0x2E)
MUTED = RGBColor(0x6B, 0x6B, 0x8A)

def set_font(run, size, bold=False, color=None, italic=False):
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def add_paragraph(doc, text, size=11, bold=False, color=None, italic=False,
                  space_before=0, space_after=6, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_font(run, size, bold, color, italic)
    return p

def add_section_heading(doc, number, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    run_num = p.add_run(str(number) + ". ")
    set_font(run_num, 13, bold=True, color=PURPLE)
    run_title = p.add_run(title)
    set_font(run_title, 13, bold=True, color=PURPLE)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(text)
    set_font(run, 10.5, color=DARK)
    return p

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D2D6EF')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ── COVER ─────────────────────────────────────────────────────────────────────

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(40)
p.paragraph_format.space_after = Pt(0)
run = p.add_run("saferspaces")
set_font(run, 28, bold=True, color=PURPLE)

p2 = doc.add_paragraph()
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after = Pt(0)
run2 = p2.add_run("GmbH")
set_font(run2, 14, color=MUTED)

add_paragraph(doc, "", space_before=40, space_after=0)

p3 = doc.add_paragraph()
p3.paragraph_format.space_before = Pt(0)
p3.paragraph_format.space_after = Pt(8)
run3 = p3.add_run("Code of Conduct")
set_font(run3, 32, bold=True, color=DARK)

add_paragraph(doc, "Verbindliche Grundsätze fuer unser Team und unsere Zusammenarbeit",
              size=12, color=MUTED, space_before=0, space_after=60)

add_divider(doc)

add_paragraph(doc, "Stand: Maerz 2026  |  Version 1.0", size=9.5, color=MUTED,
              space_before=6, space_after=4)

doc.add_page_break()

# ── PRAAMBEL ──────────────────────────────────────────────────────────────────

add_paragraph(doc, "Präambel", size=16, bold=True, color=PURPLE, space_before=0, space_after=8)

add_paragraph(doc,
    "saferspaces steht für mehr als digitale Infrastruktur. Wir sind ein Team, das der "
    "Wunsch antreibt, wirklich etwas zu verändern - professionell, tiefgehend und nachhaltig.",
    size=11, space_before=0, space_after=8)

add_paragraph(doc,
    "Wir kommen aus den Bereichen Projektmanagement, Technologie, Design und Psychologie. "
    "Eigene Erfahrungen und Beobachtungen in unseren Branchen haben uns zur gemeinsamen "
    "Überzeugung gefuehrt: Das Melden von Übergriffen und Diskriminierungen muss einfacher "
    "werden. Betroffenen muss aktiver Schutz geboten werden. Und es braucht Daten, die einen "
    "notwendigen strukturellen Wandel unterstützen.",
    size=11, space_before=0, space_after=8)

add_paragraph(doc,
    "Dieser Code of Conduct legt fest, wie wir miteinander und mit unseren Kundinnen und Kunden, "
    "Partnerorganisationen und der Öffentlichkeit umgehen. Er gilt für alle Personen, die fuer "
    "oder mit saferspaces taetig sind - hauptamtlich, freiberuflich oder ehrenamtlich.",
    size=11, space_before=0, space_after=8)

# Quote block
p_quote = doc.add_paragraph()
p_quote.paragraph_format.space_before = Pt(12)
p_quote.paragraph_format.space_after = Pt(12)
p_quote.paragraph_format.left_indent = Cm(1.2)
pPr = p_quote._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
left_bdr = OxmlElement('w:left')
left_bdr.set(qn('w:val'), 'single')
left_bdr.set(qn('w:sz'), '8')
left_bdr.set(qn('w:space'), '12')
left_bdr.set(qn('w:color'), '2A1482')
pBdr.append(left_bdr)
pPr.append(pBdr)
run_q = p_quote.add_run(
    "Die Scham muss die Seite wechseln - das ist kein Slogan, das ist eine Tatsache, "
    "die wir mit saferspaces unterstützen.")
set_font(run_q, 11, italic=True, color=PURPLE)

add_divider(doc)

# ── 1 GELTUNGSBEREICH ─────────────────────────────────────────────────────────
add_section_heading(doc, 1, "Geltungsbereich")
add_paragraph(doc,
    "Dieser Code of Conduct gilt für alle Personen, die im Namen oder Auftrag von saferspaces "
    "handeln. Dazu gehören:",
    size=11, space_before=4, space_after=4)
add_bullet(doc, "Festangestellte Mitarbeitende")
add_bullet(doc, "Freiberufliche und projektbezogen tätige Personen")
add_bullet(doc, "Partnerorganisationen und deren Personal, soweit sie saferspaces vertreten")
add_bullet(doc, "Praktikantinnen und Praktikanten")
add_paragraph(doc,
    "Er gilt fuer sämtliche beruflichen Kontexte: im Büro, bei Kundenterminen, auf "
    "Veranstaltungen, in digitaler Kommunikation und in öffentlichen Auftritten.",
    size=11, space_before=6, space_after=4)
add_divider(doc)

# ── 2 WERTE ───────────────────────────────────────────────────────────────────
add_section_heading(doc, 2, "Unsere Werte")
add_paragraph(doc,
    "Die Arbeit von saferspaces basiert auf folgenden Grundwerten, die unser Handeln leiten:",
    size=11, space_before=4, space_after=4)

values = [
    ("Schutz von Betroffenen",
     "Wir stellen die Sicherheit und das Wohlbefinden betroffener Personen in den Mittelpunkt "
     "jeder Entscheidung."),
    ("Anonymität & Vertrauen",
     "Wir erheben keine personenbezogenen Daten, keine Standortdaten und keine "
     "Gerätekennungen. Anonymität ist bei uns technisch sichergestellt - nicht nur versprochen."),
    ("Haltung zeigen",
     "Wir nehmen klar Stellung gegen Diskriminierung, Belaestigung und Ausgrenzung - "
     "intern wie extern."),
    ("Professionalität",
     "Wir handeln verantwortungsvoll, zuverlässig und mit fachlicher Sorgfalt."),
    ("Transparenz",
     "Wir kommunizieren offen - gegenueber unserem Team, unserer Kundschaft und der "
     "Öffentlichkeit."),
    ("Nachhaltigkeit & Wirkung",
     "Wir bauen auf langfristige Wirkung statt kurzfristiges Wachstum."),
]
for title, desc in values:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(title + ": ")
    set_font(r1, 10.5, bold=True, color=PURPLE)
    r2 = p.add_run(desc)
    set_font(r2, 10.5, color=DARK)
add_divider(doc)

# ── 3 TEAM ────────────────────────────────────────────────────────────────────
add_section_heading(doc, 3, "Verhalten im Team")
add_paragraph(doc,
    "Wir schaffen ein Arbeitsumfeld, das von gegenseitigem Respekt, psychologischer Sicherheit "
    "und konstruktiver Zusammenarbeit gepraegt ist. Konkret bedeutet das:",
    size=11, space_before=4, space_after=4)
add_bullet(doc, "Wir hören einander zu und nehmen unterschiedliche Perspektiven ernst.")
add_bullet(doc, "Wir kommunizieren wertschaetzend - auch in schwierigen Situationen.")
add_bullet(doc, "Wir gehen konstruktiv mit Fehlern um und nutzen sie als Lernquelle.")
add_bullet(doc, "Wir unterstützen uns gegenseitig und teilen Wissen aktiv.")
add_bullet(doc, "Wir sprechen Probleme direkt und zeitnah an, anstatt sie zu ignorieren.")
add_paragraph(doc,
    "Nicht toleriert werden: Diskriminierung, Mobbing, Belaestigung, Ausgrenzung oder jede "
    "Form von abwertendem Verhalten - unabhängig von Hierarchieebene oder Anstellungsverhältnis.",
    size=11, space_before=6, space_after=4)
add_divider(doc)

# ── 4 KUNDSCHAFT ──────────────────────────────────────────────────────────────
add_section_heading(doc, 4, "Umgang mit Kundschaft und Partnerorganisationen")
add_paragraph(doc,
    "In der Zusammenarbeit mit Kundinnen, Kunden und Partnerorganisationen verpflichten wir uns zu:",
    size=11, space_before=4, space_after=4)
add_bullet(doc, "Ehrlichkeit ueber Leistungsgrenzen und Einsatzmoeglichkeiten des Systems")
add_bullet(doc, "Diskretion im Umgang mit internen Informationen und Daten")
add_bullet(doc, "Verlässlichkeit bei vereinbarten Leistungen und Fristen")
add_bullet(doc, "Respektvollem Umgang, unabhängig von Unternehmensgröße oder Budget")
add_bullet(doc, "Klarer Kommunikation ueber Verantwortlichkeiten und Zustaendigkeiten")
add_paragraph(doc,
    "Wir lehnen Zusammenarbeiten ab, die unseren Werten widersprechen oder das Ansehen "
    "von saferspaces schädigen könnten.",
    size=11, space_before=6, space_after=4)
add_divider(doc)

# ── 5 DATENSCHUTZ ─────────────────────────────────────────────────────────────
add_section_heading(doc, 5, "Datenschutz und Vertraulichkeit")
add_paragraph(doc,
    "Datenschutz ist bei saferspaces keine Pflichübung, sondern Teil unserer Grundhaltung. "
    "Wir sind nach dem Prinzip DSGVO by Design entwickelt: Wir erheben keine personenbezogenen "
    "Daten, keine Standortdaten und keine Gerätekennungen.",
    size=11, space_before=4, space_after=6)
add_bullet(doc, "Vertrauliche Informationen aus Kundenprojekten werden nicht weitergegeben.")
add_bullet(doc, "Zugangsdaten und sicherheitsrelevante Informationen werden sicher verwahrt.")
add_bullet(doc, "Datenpannen oder -lecks werden unverzueglich intern gemeldet und dokumentiert.")
add_bullet(doc, "Unser Hosting erfolgt auf deutschen und europaeischen Servern.")
add_divider(doc)

# ── 6 INTERESSENKONFLIKTE ─────────────────────────────────────────────────────
add_section_heading(doc, 6, "Interessenkonflikte")
add_paragraph(doc,
    "Alle Teammitglieder sind verpflichtet, potenzielle Interessenkonflikte offen zu legen. "
    "Ein Interessenkonflikt liegt insbesondere vor, wenn:",
    size=11, space_before=4, space_after=4)
add_bullet(doc, "private oder finanzielle Interessen mit denen von saferspaces in Konflikt stehen,")
add_bullet(doc, "Nebentätigkeiten die Arbeit für saferspaces beeinträchtigen könnten,")
add_bullet(doc, "persönliche Beziehungen Geschäftsentscheidungen beeinflussen könnten.")
add_paragraph(doc,
    "Im Zweifel gilt: offen ansprechen, bevor eine Entscheidung getroffen wird.",
    size=11, space_before=6, space_after=4)
add_divider(doc)

# ── 7 OEFFENTLICHKEIT ─────────────────────────────────────────────────────────
add_section_heading(doc, 7, "Öffentlichkeit und Außenauftritt")
add_paragraph(doc,
    "Wer oeffentlich für saferspaces auftritt - auf Veranstaltungen, in sozialen Medien "
    "oder in Medienberichten - repraesentiert das Unternehmen und seine Werte. Wir erwarten:",
    size=11, space_before=4, space_after=4)
add_bullet(doc, "Konsistenz zwischen gelebten Werten und öffentlichen Aussagen")
add_bullet(doc, "Keine nicht abgestimmten öffentlichen Stellungnahmen im Namen von saferspaces")
add_bullet(doc, "Sensiblen Umgang mit Themen rund um Belaestigung, Diskriminierung und Betroffene")
add_bullet(doc, "Keine Weitergabe von Interna oder vertraulichen Informationen")
add_divider(doc)

# ── 8 MELDUNG ─────────────────────────────────────────────────────────────────
add_section_heading(doc, 8, "Meldung von Verstößen")
add_paragraph(doc,
    "Verstoesse gegen diesen Code of Conduct koennen jederzeit gemeldet werden – intern und "
    "vertraulich. Wir garantieren, dass Meldungen ohne negative Konsequenzen fuer die "
    "meldende Person entgegengenommen werden.",
    size=11, space_before=4, space_after=6)
add_paragraph(doc,
    "Meldungen können an die Geschaeftsfuehrung oder an eine designierte Vertrauensperson "
    "gerichtet werden. Alle Meldungen werden ernst genommen, geprüft und vertraulich behandelt.",
    size=11, space_before=0, space_after=6)
add_paragraph(doc,
    "Bei schwerwiegenden Verstößen behält sich saferspaces arbeitsrechtliche oder "
    "vertragsrechtliche Konsequenzen vor.",
    size=11, space_before=0, space_after=4)
add_divider(doc)

# ── 9 INKRAFTTRETEN ───────────────────────────────────────────────────────────
add_section_heading(doc, 9, "Inkrafttreten und Aktualisierung")
add_paragraph(doc,
    "Dieser Code of Conduct tritt mit dem Datum auf dem Deckblatt in Kraft. Er wird "
    "regelmäßig überprüft und bei Bedarf aktualisiert. Alle Teammitglieder werden "
    "ueber Änderungen informiert.",
    size=11, space_before=4, space_after=6)
add_paragraph(doc,
    "Mit der Aufnahme der Tätigkeit für saferspaces - in welcher Form auch immer - "
    "erkennen alle beteiligten Personen diesen Code of Conduct an.",
    size=11, space_before=0, space_after=4)
add_divider(doc)

add_paragraph(doc, "", space_before=10, space_after=0)
add_paragraph(doc,
    "Für Fragen oder Rückmeldungen zu diesem Dokument: hallo@saferspaces.de",
    size=10, color=MUTED, space_before=0, space_after=0, italic=True)

output_path = "/Users/leahrott/Developer/Website v2/saferspaces-code-of-conduct.docx"
doc.save(output_path)
print("Gespeichert: " + output_path)
